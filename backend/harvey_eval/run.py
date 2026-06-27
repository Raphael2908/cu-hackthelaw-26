"""Run our supervision-cockpit pipeline against a Harvey Labs task.

This is the Track A + Track B runner. It ingests a Harvey task (instructions +
documents), drives our real worker -> checker -> ranker pipeline over it, renders
the worker's output as the expected `.docx` deliverable (so Harvey's LLM judge can
grade it), and records our supervision signals (the three checker signals + the
ranker's lane/priority) for the Track B join.

It deliberately reuses the *shipped* services (app.services.worker / checker /
ranker) rather than re-implementing them, so the number reflects our system, not a
bespoke agent. See system-design/harvey-eval.md for the design + the known
worker<->task impedance mismatch (firm-standard / CELEX citation coupling).

Usage (from backend/, env injected from the repo-root .env):
    uv run python -m harvey_eval.run --task antitrust-competition/<slug> --severity high
"""

from __future__ import annotations

import argparse
import json
import os
from datetime import datetime, timezone
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
HARVEY_ROOT = REPO_ROOT / "harvey-labs"
TASKS_ROOT = HARVEY_ROOT / "tasks"
RESULTS_ROOT = HARVEY_ROOT / "results"


def _load_root_env() -> None:
    """Inject the repo-root .env (PROVIDER_MODE=real, ANTHROPIC_API_KEY, ...) into
    the environment *before* app.config.settings is imported. The backend config
    looks for .env relative to the backend dir, so running from backend/ otherwise
    silently falls back to mock mode."""
    env_path = REPO_ROOT / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        # strip inline comments + quotes
        value = value.split("#", 1)[0].strip().strip('"').strip("'")
        key = key.strip()
        if key:
            os.environ[key] = value


_load_root_env()

# app.* imports must come after env injection so settings reads real mode.
from app.db.repo import InMemoryRepo  # noqa: E402
from app.db.tables import CORPUS, FLAGS, RISK_SCORES  # noqa: E402
from app.providers.factory import get_llm_provider  # noqa: E402
from app.services.checker import run_checks  # noqa: E402
from app.services.ranker import score_task  # noqa: E402
from app.services.worker import run_review  # noqa: E402


# ── Task loading + document extraction ────────────────────────────────────────

def load_task(task_id: str) -> dict:
    task_dir = TASKS_ROOT / Path(*task_id.split("/"))
    config = json.loads((task_dir / "task.json").read_text(encoding="utf-8"))
    instructions = config.get("instructions") or (task_dir / "instructions.md").read_text("utf-8")
    return {"id": task_id, "dir": task_dir, "config": config, "instructions": instructions}


def _extract(path: Path) -> str:
    """Plain text for a task document. Reuses the backend extractor for pdf/docx/
    text; falls back to a best-effort utf-8 decode for anything else (e.g. .eml)."""
    from app.services.documents import extract_text

    raw = path.read_bytes()
    try:
        return extract_text(path.name, raw)
    except ValueError:
        return raw.decode("utf-8", errors="ignore").strip()


def build_bundle(task: dict) -> tuple[str, list[str]]:
    """Concatenate every task document into a single labelled 'draft bundle' so the
    worker sees the full context the rubric references (market memos, emails, etc.),
    not just the primary contract. Returns (bundle_text, included_filenames)."""
    docs_dir = task["dir"] / "documents"
    parts, names = [], []
    for path in sorted(docs_dir.rglob("*")):
        if not path.is_file():
            continue
        text = _extract(path)
        if not text:
            continue
        parts.append(f"===== DOCUMENT: {path.name} =====\n{text}")
        names.append(path.name)
    return "\n\n".join(parts), names


def deliverable_name(config: dict) -> str:
    """The expected output filename. Exact-naming it lets Harvey's deliverable
    matcher short-circuit (no fuzzy/LLM fallback)."""
    deliv = config.get("deliverables")
    if isinstance(deliv, dict) and deliv:
        return next(iter(deliv))
    if isinstance(deliv, list) and deliv:
        return deliv[0]
    # fall back to the most common criterion-level deliverable
    from collections import Counter

    counter: Counter[str] = Counter()
    for c in config.get("criteria", []):
        for d in c.get("deliverables", []):
            counter[d] += 1
    if counter:
        return counter.most_common(1)[0][0]
    return "deliverable.docx"


# ── Deliverable rendering ─────────────────────────────────────────────────────

def render_docx(submission: dict, task: dict, out_path: Path) -> None:
    """Render the worker's structured submission as an advisory memo .docx — the
    artifact the judge grades. The worker emits checkable findings, never a verdict
    (architecture.md §14.1); the memo preserves that."""
    import docx

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document = docx.Document()
    document.add_heading(task["config"].get("title", task["id"]), level=0)

    document.add_heading("Summary", level=1)
    document.add_paragraph(submission.get("summary", ""))

    document.add_heading("Findings", level=1)
    for i, f in enumerate(submission.get("findings", []), 1):
        clause = f.get("clause_ref") or f.get("id") or f"Finding {i}"
        document.add_heading(f"{i}. {clause}", level=2)
        document.add_paragraph(f.get("statement", ""))
        cit = f.get("citation")
        if cit:
            ref = cit.get("celex") or cit.get("source") or ""
            claim = cit.get("claim", "")
            document.add_paragraph(f"Citation: {ref} — {claim}").italic = True

    relied = submission.get("clauses_relied_on") or []
    if relied:
        document.add_heading("Clauses relied on", level=1)
        for c in relied:
            document.add_paragraph(str(c), style="List Bullet")

    document.save(str(out_path))


# ── Main ──────────────────────────────────────────────────────────────────────

def run(task_id: str, severity: str, run_id: str | None) -> dict:
    task = load_task(task_id)
    run_id = run_id or f"{task_id}/our-pipeline/{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    run_dir = RESULTS_ROOT / run_id
    out_dir = run_dir / "output"

    bundle_text, doc_names = build_bundle(task)
    provider = get_llm_provider()
    repo = InMemoryRepo()

    case_id = "harvey-case"
    # The whole document bundle is the 'draft' the worker reviews.
    bundle = repo.insert(CORPUS, {
        "kind": "draft", "case_id": case_id,
        "title": task["config"].get("title", task_id),
        "text": bundle_text, "celex": None,
    })
    # A synthetic firm standard = the task instructions (Harvey has no firm standard;
    # this is the closest faithful mapping — see the mismatch note in the design doc).
    std = repo.insert(CORPUS, {
        "kind": "firm_standard", "case_id": case_id,
        "title": "Task instructions (as standard)",
        "text": task["instructions"], "clauses": {}, "celex": None,
    })

    wtask = {
        "id": "harvey-task", "case_id": case_id,
        "target_document_id": bundle["id"], "firm_standard_id": std["id"],
        "input_process_section": task["instructions"], "severity": severity,
    }

    # Track A: worker produces the deliverable.
    submission = run_review(repo, task=wtask, provider=provider, produced_by="ai")
    deliv = deliverable_name(task["config"])
    render_docx(submission, task, out_dir / deliv)

    # Track B: checker + ranker over the same output → supervision signals.
    signals = run_checks(repo, wtask, submission, provider)
    risk = score_task(repo, wtask, signals)

    # Minimal metrics so the judge report has something; our signal payload alongside.
    (run_dir / "metrics.json").write_text(json.dumps({
        "task": task_id, "run_id": run_id, "model": os.environ.get("ANTHROPIC_MODEL", ""),
        "documents_included": doc_names, "deliverable": deliv,
    }, indent=2))

    our = {
        "task": task_id, "run_id": run_id, "severity": severity,
        "deliverable": deliv, "documents_included": doc_names,
        "n_findings": len(submission.get("findings", [])),
        "citation_support_rate": signals["citation_support_rate"],
        "deviation_score": signals["deviation_score"],
        "disagreement_score": signals["disagreement_score"],
        "uncertainty": risk["uncertainty"],
        "priority": risk["priority"],
        "lane": risk["lane"],
        "sampled": risk["sampled"],
        "has_hard_flag": signals["has_hard_flag"],
        "n_flags": len(repo.list(FLAGS)),
        "flags": [
            {"signal_type": f["signal_type"], "hard": f.get("hard", False), "title": f["title"]}
            for f in repo.list(FLAGS)
        ],
        "_risk_scores_rows": len(repo.list(RISK_SCORES)),
    }
    (run_dir / "our_eval.json").write_text(json.dumps(our, indent=2))
    print(json.dumps({k: v for k, v in our.items() if k != "flags"}, indent=2))
    print(f"\nDeliverable: {out_dir / deliv}")
    print(f"Run dir:     {run_dir}")
    print(f"Score next:  uv run --project ../harvey-labs python -m evaluation.run_eval "
          f"--run-id '{run_id}' --task '{task_id}'")
    return our


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--task", required=True)
    p.add_argument("--severity", default="high", choices=["low", "medium", "high", "extreme"])
    p.add_argument("--run-id", default=None)
    args = p.parse_args()
    run(args.task, args.severity, args.run_id)
